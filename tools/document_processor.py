#!/usr/bin/env python3
"""
Copyright (c) 2024 Theseus AI Technologies. All rights reserved.

This software is proprietary and confidential. Commercial use requires 
a valid enterprise license from Theseus AI Technologies.

Enterprise Contact: enterprise@theseus-ai.com
Website: https://theseus-ai.com/enterprise
"""

"""
Enhanced Document Processor für mlx-langchain-lite
Nutzt Batch-Embedding Engine für optimierte Verarbeitung
"""

import asyncio
import hashlib
import json
import logging
import mimetypes
import re
import time
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Union
import uuid
from concurrent.futures import ThreadPoolExecutor

import fitz  # PyMuPDF
import mammoth
import pandas as pd
from docx import Document
from markdown import markdown
from bs4 import BeautifulSoup
import numpy as np

from mlx_components.embedding_engine import EmbeddingEngine
from mlx_components.vector_store import VectorStore

logger = logging.getLogger(__name__)

@dataclass
class DocumentConfig:
    """Enhanced configuration for document processing"""
    chunk_size: int = 512
    chunk_overlap: int = 50
    max_file_size_mb: int = 50
    supported_formats: List[str] = None
    batch_size: int = 32
    enable_ocr: bool = False
    enable_table_extraction: bool = True
    enable_metadata_extraction: bool = True
    enable_pii_filtering: bool = True
    language_detection: bool = True
    quality_threshold: float = 0.7
    
    def __post_init__(self):
        if self.supported_formats is None:
            self.supported_formats = [
                'pdf', 'docx', 'doc', 'txt', 'md', 'html', 
                'csv', 'xlsx', 'xls', 'json', 'xml'
            ]

@dataclass
class DocumentChunk:
    """Enhanced document chunk with metadata"""
    content: str
    chunk_id: str
    document_id: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]
    embedding: Optional[np.ndarray] = None
    quality_score: float = 1.0
    language: Optional[str] = None
    contains_pii: bool = False
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for storage"""
        data = asdict(self)
        if self.embedding is not None:
            data['embedding'] = self.embedding.tolist()
        return data

@dataclass
class ProcessingResult:
    """Result of document processing"""
    document_id: str
    filename: str
    file_type: str
    total_chunks: int
    processing_time: float
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    errors: List[str] = None
    
    def __post_init__(self):
        if self.errors is None:
            self.errors = []

class PIIFilter:
    """PII (Personally Identifiable Information) detection and filtering"""
    
    # Patterns for common PII
    PATTERNS = {
        'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        'phone': r'(\+\d{1,3}[-.\s]?)?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
        'ssn': r'\b\d{3}-\d{2}-\d{4}\b',
        'credit_card': r'\b\d{4}[-.\s]?\d{4}[-.\s]?\d{4}[-.\s]?\d{4}\b',
        'ip_address': r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b',
    }
    
    @classmethod
    def detect_pii(cls, text: str) -> Dict[str, List[str]]:
        """Detect PII in text"""
        detected = {}
        
        for pii_type, pattern in cls.PATTERNS.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                detected[pii_type] = matches
        
        return detected
    
    @classmethod
    def filter_pii(cls, text: str, replacement: str = "[REDACTED]") -> str:
        """Remove PII from text"""
        filtered_text = text
        
        for pii_type, pattern in cls.PATTERNS.items():
            filtered_text = re.sub(pattern, replacement, filtered_text, flags=re.IGNORECASE)
        
        return filtered_text

class LanguageDetector:
    """Simple language detection"""
    
    # Common words for language detection
    LANGUAGE_INDICATORS = {
        'en': ['the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 'how', 'man', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy', 'did', 'its', 'let', 'put', 'say', 'she', 'too', 'use'],
        'de': ['der', 'die', 'und', 'in', 'den', 'von', 'zu', 'das', 'mit', 'sich', 'des', 'auf', 'für', 'ist', 'im', 'dem', 'nicht', 'ein', 'eine', 'als', 'auch', 'es', 'an', 'werden', 'aus', 'er', 'hat', 'dass', 'sie', 'nach', 'wird', 'bei', 'einer', 'um', 'am', 'sind', 'noch', 'wie', 'einem', 'über', 'einen', 'so', 'zum', 'war', 'haben', 'nur', 'oder', 'aber', 'vor', 'zur', 'bis', 'mehr', 'durch', 'man', 'sein', 'wurde', 'sei', 'in'],
        'fr': ['le', 'de', 'et', 'à', 'un', 'il', 'être', 'et', 'en', 'avoir', 'que', 'pour', 'dans', 'ce', 'son', 'une', 'sur', 'avec', 'ne', 'se', 'pas', 'tout', 'plus', 'par', 'grand', 'il', 'me', 'même', 'faire', 'elle', 'si', 'lors', 'mon', 'man', 'qui', 'lui', 'nous', 'comme', 'mais', 'pouvoir', 'eux', 'très', 'lors', 'sans', 'can', 'lieu', 'où', 'encore', 'aussi', 'alors'],
        'es': ['el', 'la', 'de', 'que', 'y', 'a', 'en', 'un', 'ser', 'se', 'no', 'te', 'lo', 'le', 'da', 'su', 'por', 'son', 'con', 'para', 'mi', 'está', 'si', 'pero', 'me', 'one', 'bien', 'has', 'ese', 'va', 'ya', 'todo', 'voy', 'muy', 'hay', 'ahora', 'algo', 'estoy', 'tengo', 'nos', 'tú', 'nada', 'cuando', 'ha', 'este', 'sí', 'antes', 'tiempo', 'después', 'oye', 'aunque', 'donde', 'todavía', 'pueden', 'hasta']
    }
    
    @classmethod
    def detect_language(cls, text: str) -> str:
        """Detect language of text"""
        if not text or len(text.strip()) < 10:
            return 'unknown'
        
        words = text.lower().split()[:100]  # Check first 100 words
        
        scores = {}
        for lang, indicators in cls.LANGUAGE_INDICATORS.items():
            score = sum(1 for word in words if word in indicators)
            scores[lang] = score / len(words) if words else 0
        
        if not scores:
            return 'unknown'
        
        detected_lang = max(scores, key=scores.get)
        return detected_lang if scores[detected_lang] > 0.1 else 'unknown'

class ChunkingStrategy:
    """Advanced chunking strategies"""
    
    @staticmethod
    def semantic_chunking(text: str, chunk_size: int, overlap: int) -> List[Tuple[str, int, int]]:
        """Semantic-aware chunking"""
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        current_start = 0
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) <= chunk_size:
                current_chunk += paragraph + "\n\n"
            else:
                if current_chunk.strip():
                    end_pos = current_start + len(current_chunk)
                    chunks.append((current_chunk.strip(), current_start, end_pos))
                    
                    # Handle overlap
                    if overlap > 0:
                        overlap_start = max(0, end_pos - overlap)
                        current_start = overlap_start
                        current_chunk = text[overlap_start:end_pos] + paragraph + "\n\n"
                    else:
                        current_start = end_pos
                        current_chunk = paragraph + "\n\n"
                else:
                    current_chunk = paragraph + "\n\n"
        
        # Add remaining chunk
        if current_chunk.strip():
            end_pos = current_start + len(current_chunk)
            chunks.append((current_chunk.strip(), current_start, end_pos))
        
        return chunks
    
    @staticmethod
    def sentence_aware_chunking(text: str, chunk_size: int, overlap: int) -> List[Tuple[str, int, int]]:
        """Sentence-boundary aware chunking"""
        # Simple sentence splitting
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = ""
        current_start = 0
        sentence_start = 0
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += sentence + " "
            else:
                if current_chunk.strip():
                    end_pos = sentence_start
                    chunks.append((current_chunk.strip(), current_start, end_pos))
                    
                    # Handle overlap
                    if overlap > 0:
                        overlap_start = max(0, end_pos - overlap)
                        current_start = overlap_start
                        overlap_text = text[overlap_start:end_pos]
                        current_chunk = overlap_text + sentence + " "
                    else:
                        current_start = end_pos
                        current_chunk = sentence + " "
                else:
                    current_chunk = sentence + " "
            
            sentence_start += len(sentence) + 1
        
        # Add remaining chunk
        if current_chunk.strip():
            chunks.append((current_chunk.strip(), current_start, len(text)))
        
        return chunks

class DocumentProcessor:
    """Enhanced document processor with batch embedding"""
    
    def __init__(self, config: DocumentConfig, embedding_engine: EmbeddingEngine, vector_store: VectorStore):
        self.config = config
        self.embedding_engine = embedding_engine
        self.vector_store = vector_store
        self.executor = ThreadPoolExecutor(max_workers=4)
        
        # File type processors mapping
        self.processors = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'doc': self._process_doc,
            'txt': self._process_text,
            'md': self._process_markdown,
            'html': self._process_html,
            'csv': self._process_csv,
            'xlsx': self._process_excel,
            'xls': self._process_excel,
            'json': self._process_json,
            'xml': self._process_xml
        }
    
    async def process_document(self, file_path: str, user_id: str, 
                             document_type: Optional[str] = None) -> ProcessingResult:
        """Process single document with batch embedding"""
        start_time = time.time()
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type
        if document_type is None:
            document_type = self._detect_file_type(file_path)
        
        if document_type not in self.config.supported_formats:
            raise ValueError(f"Unsupported file type: {document_type}")
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.config.max_file_size_mb:
            raise ValueError(f"File too large: {file_size_mb:.1f}MB > {self.config.max_file_size_mb}MB")
        
        document_id = str(uuid.uuid4())
        
        try:
            # Extract text and metadata
            content, metadata = await self._extract_content(file_path, document_type)
            
            # Create chunks
            chunks = await self._create_chunks(content, document_id, metadata)
            
            # Filter PII if enabled
            if self.config.enable_pii_filtering:
                chunks = self._filter_pii_from_chunks(chunks)
            
            # Detect language if enabled
            if self.config.language_detection:
                chunks = self._detect_language_in_chunks(chunks)
            
            # Calculate quality scores
            chunks = self._calculate_quality_scores(chunks)
            
            # Filter by quality threshold
            high_quality_chunks = [c for c in chunks if c.quality_score >= self.config.quality_threshold]
            
            # Create embeddings in batches
            chunks_with_embeddings = await self._create_batch_embeddings(high_quality_chunks)
            
            # Store in vector database
            await self._store_chunks(chunks_with_embeddings, user_id)
            
            processing_time = time.time() - start_time
            
            result = ProcessingResult(
                document_id=document_id,
                filename=file_path.name,
                file_type=document_type,
                total_chunks=len(chunks_with_embeddings),
                processing_time=processing_time,
                chunks=chunks_with_embeddings,
                metadata=metadata
            )
            
            logger.info(f"Processed document {file_path.name}: {len(chunks_with_embeddings)} chunks in {processing_time:.2f}s")
            return result
            
        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"Error processing document {file_path.name}: {e}")
            
            return ProcessingResult(
                document_id=document_id,
                filename=file_path.name,
                file_type=document_type,
                total_chunks=0,
                processing_time=processing_time,
                chunks=[],
                metadata={},
                errors=[str(e)]
            )
    
    async def process_documents_batch(self, file_paths: List[str], user_id: str) -> List[ProcessingResult]:
        """Process multiple documents in batch"""
        logger.info(f"Processing batch of {len(file_paths)} documents")
        
        # Process documents in parallel
        tasks = []
        for file_path in file_paths:
            task = self.process_document(file_path, user_id)
            tasks.append(task)
        
        # Execute in batches to avoid overwhelming the system
        batch_size = 5
        results = []
        
        for i in range(0, len(tasks), batch_size):
            batch = tasks[i:i + batch_size]
            batch_results = await asyncio.gather(*batch, return_exceptions=True)
            
            for result in batch_results:
                if isinstance(result, Exception):
                    logger.error(f"Batch processing error: {result}")
                    # Create error result
                    error_result = ProcessingResult(
                        document_id=str(uuid.uuid4()),
                        filename="unknown",
                        file_type="unknown",
                        total_chunks=0,
                        processing_time=0.0,
                        chunks=[],
                        metadata={},
                        errors=[str(result)]
                    )
                    results.append(error_result)
                else:
                    results.append(result)
        
        return results
    
    def _detect_file_type(self, file_path: Path) -> str:
        """Detect file type from extension and MIME type"""
        extension = file_path.suffix.lower().lstrip('.')
        
        # Direct mapping
        if extension in self.config.supported_formats:
            return extension
        
        # MIME type fallback
        mime_type, _ = mimetypes.guess_type(str(file_path))
        mime_mapping = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'text/plain': 'txt',
            'text/markdown': 'md',
            'text/html': 'html',
            'text/csv': 'csv',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
            'application/vnd.ms-excel': 'xls',
            'application/json': 'json',
            'application/xml': 'xml',
            'text/xml': 'xml'
        }
        
        return mime_mapping.get(mime_type, 'txt')
    
    async def _extract_content(self, file_path: Path, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content and metadata from file"""
        processor = self.processors.get(file_type, self._process_text)
        
        # Run processor in thread pool
        loop = asyncio.get_event_loop()
        content, metadata = await loop.run_in_executor(self.executor, processor, file_path)
        
        return content, metadata
    
    def _process_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process PDF file"""
        try:
            doc = fitz.open(str(file_path))
            content = ""
            metadata = {
                'page_count': doc.page_count,
                'author': doc.metadata.get('author', ''),
                'title': doc.metadata.get('title', ''),
                'subject': doc.metadata.get('subject', ''),
                'creator': doc.metadata.get('creator', ''),
                'creation_date': doc.metadata.get('creationDate', ''),
                'modification_date': doc.metadata.get('modDate', '')
            }
            
            # Extract text from all pages
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                # Extract tables if enabled
                if self.config.enable_table_extraction:
                    tables = page.find_tables()
                    for table in tables:
                        try:
                            table_df = table.to_pandas()
                            table_text = table_df.to_string(index=False)
                            page_text += f"\n\nTable:\n{table_text}\n"
                        except:
                            pass
                
                content += f"\n--- Page {page_num + 1} ---\n{page_text}"
            
            doc.close()
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _process_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX file"""
        try:
            doc = Document(str(file_path))
            content = ""
            
            # Extract core properties
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'paragraph_count': len(doc.paragraphs)
            }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n\n"
            
            # Extract text from tables if enabled
            if self.config.enable_table_extraction:
                for table in doc.tables:
                    table_text = "Table:\n"
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        table_text += row_text + "\n"
                    content += table_text + "\n"
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _process_doc(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process DOC file using mammoth"""
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                content = result.value
                
                metadata = {
                    'extraction_warnings': result.messages,
                    'character_count': len(content)
                }
                
                return content, metadata
                
        except Exception as e:
            logger.error(f"Error processing DOC {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _process_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            metadata = {
                'character_count': len(content),
                'line_count': len(content.split('\n')),
                'encoding': 'utf-8'
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _process_markdown(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                md_content = f.read()
            
            # Convert to HTML then extract text
            html_content = markdown(md_content)
            soup = BeautifulSoup(html_content, 'html.parser')
            content = soup.get_text()
            
            metadata = {
                'markdown_length': len(md_content),
                'html_length': len(html_content),
                'text_length': len(content)
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _process_html(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract metadata
            title = soup.find('title')
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            
            metadata = {
                'title': title.get_text() if title else '',
                'description': meta_desc.get('content') if meta_desc else '',
                'html_length': len(html_content)
            }
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            
            content = soup.get_text()
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _process_csv(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process CSV file"""
        try:
            df = pd.read_csv(file_path)
            
            metadata = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'data_types': df.dtypes.to_dict()
            }
            
            # Convert to text representation
            content = f"CSV Data Summary:\nRows: {len(df)}, Columns: {len(df.columns)}\n\n"
            content += f"Column Names: {', '.join(df.columns)}\n\n"
            content += "Data Preview:\n" + df.head(10).to_string() + "\n\n"
            content += "Data Description:\n" + df.describe().to_string()
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _process_excel(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process Excel file"""
        try:
            xls = pd.ExcelFile(file_path)
            content = ""
            
            metadata = {
                'sheet_names': xls.sheet_names,
                'sheet_count': len(xls.sheet_names)
            }
            
            for sheet_name in xls.sheet_names[:5]:  # Limit to first 5 sheets
                df = pd.read_excel(xls, sheet_name=sheet_name)
                
                content += f"\n--- Sheet: {sheet_name} ---\n"
                content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                content += f"Columns: {', '.join(df.columns)}\n\n"
                content += "Data Preview:\n" + df.head(5).to_string() + "\n"
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _process_json(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            content = json.dumps(data, indent=2, ensure_ascii=False)
            
            metadata = {
                'json_size': len(content),
                'structure_type': type(data).__name__,
                'top_level_keys': list(data.keys()) if isinstance(data, dict) else []
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing JSON {file_path}: {e}")
            return "", {"error": str(e)}
    
    def _process_xml(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process XML file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                xml_content = f.read()
            
            soup = BeautifulSoup(xml_content, 'xml')
            content = soup.get_text()
            
            metadata = {
                'xml_length': len(xml_content),
                'text_length': len(content),
                'root_tag': soup.find().name if soup.find() else ''
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing XML {file_path}: {e}")
            return "", {"error": str(e)}
    
    async def _create_chunks(self, content: str, document_id: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Create chunks from content using advanced strategies"""
        if not content.strip():
            return []
        
        # Choose chunking strategy based on content type
        if len(content) > 10000:  # Large document
            chunks_data = ChunkingStrategy.semantic_chunking(
                content, self.config.chunk_size, self.config.chunk_overlap
            )
        else:  # Smaller document
            chunks_data = ChunkingStrategy.sentence_aware_chunking(
                content, self.config.chunk_size, self.config.chunk_overlap
            )
        
        chunks = []
        for i, (chunk_content, start_char, end_char) in enumerate(chunks_data):
            if chunk_content.strip():
                chunk = DocumentChunk(
                    content=chunk_content,
                    chunk_id=f"{document_id}_{i}",
                    document_id=document_id,
                    chunk_index=i,
                    start_char=start_char,
                    end_char=end_char,
                    metadata={
                        **metadata,
                        'chunk_size': len(chunk_content),
                        'word_count': len(chunk_content.split())
                    }
                )
                chunks.append(chunk)
        
        return chunks
    
    def _filter_pii_from_chunks(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Filter PII from chunks"""
        filtered_chunks = []
        
        for chunk in chunks:
            pii_detected = PIIFilter.detect_pii(chunk.content)
            
            if pii_detected:
                chunk.contains_pii = True
                chunk.metadata['pii_detected'] = pii_detected
                # Filter the content
                chunk.content = PIIFilter.filter_pii(chunk.content)
            
            filtered_chunks.append(chunk)
        
        return filtered_chunks
    
    def _detect_language_in_chunks(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Detect language in chunks"""
        for chunk in chunks:
            chunk.language = LanguageDetector.detect_language(chunk.content)
            chunk.metadata['language'] = chunk.language
        
        return chunks
    
    def _calculate_quality_scores(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Calculate quality scores for chunks"""
        for chunk in chunks:
            score = 1.0
            
            # Length penalty for very short chunks
            if len(chunk.content) < 50:
                score *= 0.5
            
            # Word count factor
            word_count = len(chunk.content.split())
            if word_count < 10:
                score *= 0.6
            elif word_count > 100:
                score *= 1.1
            
            # Special character ratio penalty
            special_chars = sum(1 for c in chunk.content if not c.isalnum() and not c.isspace())
            if special_chars / len(chunk.content) > 0.3:
                score *= 0.7
            
            # PII penalty
            if chunk.contains_pii:
                score *= 0.8
            
            # Language detection bonus
            if chunk.language and chunk.language != 'unknown':
                score *= 1.1
            
            chunk.quality_score = min(score, 1.0)
        
        return chunks
    
    async def _create_batch_embeddings(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Create embeddings for chunks in batches"""
        if not chunks:
            return chunks
        
        logger.info(f"Creating embeddings for {len(chunks)} chunks")
        
        # Prepare texts for embedding
        texts = [chunk.content for chunk in chunks]
        
        # Process in batches
        batch_size = self.config.batch_size
        embedded_chunks = []
        
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i + batch_size]
            batch_chunks = chunks[i:i + batch_size]
            
            try:
                # Generate embeddings
                embeddings = await self.embedding_engine.embed_batch(batch_texts)
                
                # Assign embeddings to chunks
                for j, embedding in enumerate(embeddings):
                    batch_chunks[j].embedding = embedding
                
                embedded_chunks.extend(batch_chunks)
                logger.info(f"Embedded batch {i//batch_size + 1}/{(len(texts) + batch_size - 1)//batch_size}")
                
            except Exception as e:
                logger.error(f"Error creating embeddings for batch {i}: {e}")
                # Add chunks without embeddings
                embedded_chunks.extend(batch_chunks)
        
        return embedded_chunks
    
    async def _store_chunks(self, chunks: List[DocumentChunk], user_id: str):
        """Store chunks in vector database"""
        if not chunks:
            return
        
        # Prepare data for vector store
        embeddings = []
        metadata_list = []
        
        for chunk in chunks:
            if chunk.embedding is not None:
                embeddings.append(chunk.embedding)
                
                chunk_metadata = {
                    'chunk_id': chunk.chunk_id,
                    'document_id': chunk.document_id,
                    'chunk_index': chunk.chunk_index,
                    'content': chunk.content,
                    'start_char': chunk.start_char,
                    'end_char': chunk.end_char,
                    'quality_score': chunk.quality_score,
                    'language': chunk.language,
                    'contains_pii': chunk.contains_pii,
                    'user_id': user_id,
                    'timestamp': time.time(),
                    **chunk.metadata
                }
                metadata_list.append(chunk_metadata)
        
        if embeddings:
            try:
                await self.vector_store.add_embeddings(
                    embeddings=embeddings,
                    metadata=metadata_list,
                    user_id=user_id,
                    collection_name=f"documents_{user_id}"
                )
                logger.info(f"Stored {len(embeddings)} chunks in vector database")
                
            except Exception as e:
                logger.error(f"Error storing chunks in vector database: {e}")
    
    async def search_documents(self, query: str, user_id: str, top_k: int = 10, 
                             filters: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """Search documents using semantic similarity"""
        try:
            # Generate query embedding
            query_embedding = await self.embedding_engine.embed_text(query)
            
            # Search in vector store
            results = await self.vector_store.search(
                query_embedding=query_embedding,
                user_id=user_id,
                collection_name=f"documents_{user_id}",
                top_k=top_k,
                filters=filters
            )
            
            return results
            
        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []
    
    async def get_document_chunks(self, document_id: str, user_id: str) -> List[Dict[str, Any]]:
        """Get all chunks for a specific document"""
        try:
            results = await self.vector_store.search_by_metadata(
                filters={"document_id": document_id, "user_id": user_id},
                collection_name=f"documents_{user_id}"
            )
            
            # Sort by chunk index
            results.sort(key=lambda x: x.get('metadata', {}).get('chunk_index', 0))
            return results
            
        except Exception as e:
            logger.error(f"Error retrieving document chunks: {e}")
            return []
    
    async def delete_document(self, document_id: str, user_id: str) -> bool:
        """Delete all chunks for a document"""
        try:
            # Get all chunk IDs for the document
            chunks = await self.get_document_chunks(document_id, user_id)
            chunk_ids = [chunk['metadata']['chunk_id'] for chunk in chunks]
            
            if chunk_ids:
                await self.vector_store.delete_by_ids(
                    ids=chunk_ids,
                    user_id=user_id,
                    collection_name=f"documents_{user_id}"
                )
                logger.info(f"Deleted {len(chunk_ids)} chunks for document {document_id}")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            return False
    
    async def get_document_statistics(self, user_id: str) -> Dict[str, Any]:
        """Get document processing statistics"""
        try:
            # Get all documents for user
            all_results = await self.vector_store.search_by_metadata(
                filters={"user_id": user_id},
                collection_name=f"documents_{user_id}",
                limit=10000  # Large limit to get all
            )
            
            # Calculate statistics
            total_chunks = len(all_results)
            documents = {}
            languages = {}
            file_types = {}
            quality_scores = []
            
            for result in all_results:
                metadata = result.get('metadata', {})
                doc_id = metadata.get('document_id')
                language = metadata.get('language', 'unknown')
                file_type = metadata.get('file_type', 'unknown')
                quality = metadata.get('quality_score', 0.0)
                
                if doc_id:
                    documents[doc_id] = documents.get(doc_id, 0) + 1
                
                languages[language] = languages.get(language, 0) + 1
                file_types[file_type] = file_types.get(file_type, 0) + 1
                quality_scores.append(quality)
            
            avg_quality = sum(quality_scores) / len(quality_scores) if quality_scores else 0.0
            
            return {
                'total_documents': len(documents),
                'total_chunks': total_chunks,
                'average_chunks_per_document': total_chunks / len(documents) if documents else 0,
                'languages': languages,
                'file_types': file_types,
                'average_quality_score': avg_quality,
                'documents_by_id': documents
            }
            
        except Exception as e:
            logger.error(f"Error getting document statistics: {e}")
            return {}
    
    async def cleanup(self):
        """Cleanup resources"""
        self.executor.shutdown(wait=True)
        logger.info("Document processor cleanup completed")

# Advanced document analysis tools
class DocumentAnalyzer:
    """Advanced document analysis capabilities"""
    
    @staticmethod
    def extract_key_phrases(text: str, top_k: int = 10) -> List[str]:
        """Extract key phrases from text"""
        # Simple n-gram extraction
        words = re.findall(r'\b\w+\b', text.lower())
        
        # Generate bigrams and trigrams
        bigrams = [f"{words[i]} {words[i+1]}" for i in range(len(words)-1)]
        trigrams = [f"{words[i]} {words[i+1]} {words[i+2]}" for i in range(len(words)-2)]
        
        # Count frequencies
        from collections import Counter
        bigram_counts = Counter(bigrams)
        trigram_counts = Counter(trigrams)
        
        # Get top phrases
        top_bigrams = [phrase for phrase, count in bigram_counts.most_common(top_k//2)]
        top_trigrams = [phrase for phrase, count in trigram_counts.most_common(top_k//2)]
        
        return top_bigrams + top_trigrams
    
    @staticmethod
    def calculate_readability_score(text: str) -> float:
        """Calculate readability score (simplified Flesch Reading Ease)"""
        sentences = len(re.split(r'[.!?]+', text))
        words = len(text.split())
        syllables = sum(DocumentAnalyzer._count_syllables(word) for word in text.split())
        
        if sentences == 0 or words == 0:
            return 0.0
        
        # Simplified Flesch Reading Ease
        score = 206.835 - (1.015 * (words / sentences)) - (84.6 * (syllables / words))
        return max(0.0, min(100.0, score))
    
    @staticmethod
    def _count_syllables(word: str) -> int:
        """Count syllables in a word (approximation)"""
        word = word.lower()
        vowels = "aeiouy"
        syllable_count = 0
        previous_was_vowel = False
        
        for char in word:
            is_vowel = char in vowels
            if is_vowel and not previous_was_vowel:
                syllable_count += 1
            previous_was_vowel = is_vowel
        
        # Adjust for silent e
        if word.endswith('e'):
            syllable_count -= 1
        
        # Every word has at least one syllable
        return max(1, syllable_count)
    
    @staticmethod
    def extract_entities(text: str) -> Dict[str, List[str]]:
        """Extract named entities (simple pattern-based)"""
        entities = {
            'dates': [],
            'numbers': [],
            'urls': [],
            'emails': [],
            'capitalized_words': []
        }
        
        # Date patterns
        date_patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',
            r'\d{4}-\d{2}-\d{2}',
            r'\b\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b'
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            entities['dates'].extend(matches)
        
        # Numbers
        numbers = re.findall(r'\b\d+(?:\.\d+)?\b', text)
        entities['numbers'] = numbers[:20]  # Limit to avoid too many
        
        # URLs
        urls = re.findall(r'https?://[^\s]+', text)
        entities['urls'] = urls
        
        # Emails
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        entities['emails'] = emails
        
        # Capitalized words (potential proper nouns)
        capitalized = re.findall(r'\b[A-Z][a-z]+\b', text)
        entities['capitalized_words'] = list(set(capitalized))[:30]  # Unique and limited
        
        return entities

# Export main classes
__all__ = [
    'DocumentProcessor', 
    'DocumentConfig', 
    'DocumentChunk', 
    'ProcessingResult',
    'PIIFilter',
    'LanguageDetector',
    'ChunkingStrategy',
    'DocumentAnalyzer'
]